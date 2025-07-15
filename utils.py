import os
from openpyxl import Workbook
from docx import Document
import os
from dotenv import load_dotenv
from openai import OpenAI
import base64
import base64
from io import BytesIO
import aspose.words as aw


class AssetExtractorUtils:
    def __init__(self, directory_path):
        load_dotenv()
        self.directory_path = directory_path
        self.workbook = None
        self.create_workbook()
        self.client = OpenAI(api_key=os.getenv('NOT_AI_API_KEY')) #gogreen

    def create_workbook(self):
        
        self.workbook = Workbook()

        sheets_with_headers = {
            "Fixed Extinguishing Systems": [
                "address", "business_name", "asset_type", "variant", "last_recharge_date", "location_of_cylinders", 
                "manufacturer", "model_number", "serial_number", "power", 
            ],
            "Fire Hoses": [
                "address", "business_name", "location", "length", 
                "size", "nozzle", "status", "next_ht", "notes"
            ],
            "Fire Hydrants": [
                "address", "business_name", "hydrant_number", "location", "make", "model",
                "color", "shutoff_location", "type"
            ], 
            "Backflows" : [
                "name_of_premise", "address", "asset_type", "variant", "manufacturer", "model_number", 
                "serial_number", "size", "location"
            ], 
            "Extinguishers": [
                "address", "business_name", "location", "size", "brand", "serial_number", 
                "manufacture_date", "next_service_date", "comments"
            ], 
            "Fire Pumps": [
                "address", "business_name", "asset_type", "variant", "location", "system", 
                "water_supply_source", "pump_manufacturer", "controller_manufacturer", "controller_model" 
            ],   
            "Emergency Lights": [
                "address", "business_name", "type", "variant", "location", "battery_size",
                "battery #", "battery_date", "voltage / size", "comments"
            ], 
            "Special Suppression" : [
                "address", "business_name", "asset_type", "variant", "make", "model"
            ], 
            "Alarm Systems" : [
                'address', 'ref', 'business_name', 'manufacturere', 'model_number', 
                'fire_signal_receiving_centre', 'ulc_serial_number'
            ], 
            "Alarm System Devices": [
                'system', 'asset_type', 'variant', 'zone_circuit_number' 
            ], 
            "Sprinkler Systems": [
                'address', 'business_name', 'asset_description'
            ],    
        }
       
        default_sheet = self.workbook.active
        self.workbook.remove(default_sheet)

        for sheet_name, headers in sheets_with_headers.items():
            sheet = self.workbook.create_sheet(title=sheet_name)
            sheet.append(headers)

        self.workbook.save("assets.xlsx")

    def update_workbook(self, sheet_name, data):
        
        sheet = self.workbook[sheet_name]

        for row in data:
            sheet.append(row)

        self.workbook.save("assets.xlsx")

    def get_document_text(self, file_path): 
    
        doc = Document(file_path)
        full_text = []

        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)

        for table in doc.tables:
            try:
                for row in table.rows:
                    try:
                        for cell in row.cells:
        
                            try:
                                full_text.append(cell.text)
                            except:
                                
                                continue
                    except:
                        #absolutely zero idea why this doesnt work without these try / except blocks lol
                        continue
            except:
                
                continue
            
        return '\n'.join(full_text)

    def get_document_tables(self, file_path): 
        doc = Document(file_path)

        return doc.tables
    
    def get_docx_files(self):

        docx_files = []
        try:
            for root, dirs, files in os.walk(self.directory_path):
                for file in files:
                    full_path = os.path.join(root, file)
                    if file.lower().endswith('.docx'):
                    
                        relative_path = os.path.relpath(full_path, self.directory_path)
                        docx_files.append(relative_path)
                    else:
                        os.remove(full_path)
           
            for root, dirs, files in os.walk(self.directory_path, topdown=False):
        
                if root == self.directory_path:
                    continue
                try:
                    if not os.listdir(root):
                        os.rmdir(root)
                except OSError:
                    pass
            
            return docx_files
        
        except OSError as e:
            print(f"Error reading directory {self.directory_path}: {e}")
            return []
        
    def _headers_match(self, actual_headers, target_headers):
        """Check if actual headers match target headers (allowing for some flexibility)"""
        if len(actual_headers) < len(target_headers):
            return False
    
        key_matches = 0
        for i, target in enumerate(target_headers):
            if i < len(actual_headers):
                # Normalize text for comparison
                actual_normalized = actual_headers[i].replace('\n', ' ').replace('\r', '').strip()
                target_normalized = target.replace('\n', ' ').replace('\r', '').strip()
                
                if actual_normalized.lower() == target_normalized.lower():
                    key_matches += 1
                elif target in ["A", "B", "C", "D", "E", "F", "G"] and actual_normalized == target:
                    key_matches += 1
        
        return key_matches >= 7  
    
    def find_header_row(self, table, target_text, w=False):
        """
        Find the row index that contains the target header text.
        Returns (row_index, column_index) if found, otherwise (None, None)
        """
    
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if target_text in cell.text.strip():
                    return row_idx, col_idx
        
        return None, None

    def get_extraction_method(self, text):
        """determine which method to use based unique words per each template. this randomly just stops working sometimes"""
        text_lower = text.lower()
        
        method_keywords = {
            'fixed_extinguishing_systems': [
                'inspection, testing and maintenance report for fixed extinguishing systems',
                'location of system cylinders'
            ],
            'fire_hoses': [
                'fire hose test and inspection'
            ], 
            'fire_hydrants' : [
                'fire hydrant inspection & testing'
            ], 
            'backflows': [
                'location of backflow preventer', 'double check assemblies'
            ], 
            'fire_pumps': [
                'fire pump annual performance tests', 
                'pump has a prv installed'
            ], 
            'smoke_alarms' : [
                'smoke alarm device record'
            ],  
            'emergency_lighting': [
                'unit emergency lighting test'
            ], 
            'emergency_lighting_extinguisher': [
                'unit emergency lighting /'
            ], 
            'extinguishers': [
                'extinguisher test & inspection'
            ], 
            'special_suppression': [
                'report for special fire suppression system', 
                'novec 1230'
            ], 
            'alarm_system_devices': [
                'nbc', 'provides single-stage operation'
            ], 
            'sprinkler_systems': [
                'is the fdc check valve free of leaks?'
            ]
        }
        
        for method_name, keywords in method_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                return method_name
        
        return None

    #oooh how mysterious i wonder what api it is!! 
    def api_call(self, file_path, page, prompt): 
        #needed this for 3 templates that had checkboxes that have the same properties when checked vs unchecked 
        #my only chance of figuring this one out was gpt, so if you think about it im really just cutting out the middle man

        img = self.doc_to_base64(file_path, page)
        
        try:
        
            response = self.client.chat.completions.create(
                model="gpt-4o",  
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{img}"
                                },
                            },
                        ],
                    }
                ],
                max_tokens=5000,
                timeout=30  
            )
            print(response.choices[0].message.content)
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"Error occurred: {e}")
            print(f"Error type: {type(e)}")
    

    def doc_to_base64(self, doc_path, page):
        try:
            doc = aw.Document(doc_path)
            options = aw.saving.ImageSaveOptions(aw.SaveFormat.PNG)
            options.horizontal_resolution = 600
            options.vertical_resolution = 600
            options.page_set = aw.saving.PageSet(page)
            
            stream = BytesIO()
            doc.save(stream, options)
            stream.seek(0)
            
            img_bytes = stream.getvalue()
            if len(img_bytes) == 0:
                return None
                
            base64_str = base64.b64encode(img_bytes).decode()
                
            return base64_str
            
        except Exception as e:
            return None

    
