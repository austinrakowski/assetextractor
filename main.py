from docx import Document
import os
import pprint

class FireProtectionReportExtractor:
    def __init__(self, directory_path):
        self.directory_path = directory_path
        
    def extract_fixed_extinguishing_systems(self, doc):
        """Extract data from Fixed Extinguishing Systems template"""
        if not doc.tables:
            return {}
            
        table = doc.tables[0]
        data = {}
        
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            
            for i in range(len(cells) - 1):
                if cells[i].endswith(':'):
                    key = cells[i].replace(':', '')

                    for j in range(i + 1, len(cells)):
                        if cells[j] != cells[i] and not cells[j].endswith(':'):
                            data[key] = cells[j]
                            break
        
        return data

    def extract_unit_emergency_lighting_extinguisher(self, doc):
        """Extract data from Unit Emergency Lighting / Extinguisher Test & Inspection template"""
        if len(doc.tables) < 3:
            return {}
        
        data_table = doc.tables[2]
        
        extinguishers = []
        headers = []
    
        first_row = data_table.rows[0]
        headers = [cell.text.strip() for cell in first_row.cells]
        
        for i, row in enumerate(data_table.rows):
            if i == 0:  
                continue
                
            row_data = [cell.text.strip() for cell in row.cells]
            
            extinguisher = {}
            for j, value in enumerate(row_data):
                if j < len(headers) and headers[j] and value:
                    extinguisher[headers[j]] = value
            
            if extinguisher: 
                extinguishers.append(extinguisher)
    
        return {'extinguishers': extinguishers}

    def extract_entinguishers(self, doc): 
        table = doc.tables[2]
        for i, row in enumerate(table.rows):
            if i == 0: 
                continue

            cells = row.cells
            
            # Handle cases where rows might have fewer cells
            try:
                location = cells[0].text.strip() if len(cells) > 0 else ""
                size_type = cells[1].text.strip() if len(cells) > 1 else ""
                brand = cells[2].text.strip() if len(cells) > 2 else ""
                serial_num = cells[3].text.strip() if len(cells) > 3 else ""
                mfg_date = cells[4].text.strip() if len(cells) > 4 else ""
                nsd = cells[5].text.strip() if len(cells) > 5 else ""
                comments = cells[6].text.strip() if len(cells) > 6 else ""
            except IndexError as e:
                print(f"Error processing row {i}: {e}")
                continue

            print([location, size_type, brand, serial_num, mfg_date, nsd, comments])

        return None

    
    def debug_tables(self, filename):
        """Debug method to see table structure"""
        file_path = os.path.join(self.directory_path, filename)
        doc = Document(file_path)
        
        print(f"Number of tables: {len(doc.tables)}")
        
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i}: {len(table.rows)} rows, {len(table.columns)} columns")
            for j, row in enumerate(table.rows):
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    print(f"  Row {j}: {row_text}")
    
    def test_single_file(self, filename, method_name):
        """Test a specific extraction method on a single file"""
        file_path = os.path.join(self.directory_path, filename)
        
        try:
            doc = Document(file_path)
            method = getattr(self, method_name)
            result = method(doc)
            print(f"File: {filename}")
            print(f"Method: {method_name}")
            pprint.pprint(result)
            return result
        except Exception as e:
            print(f"Error testing {filename} with {method_name}: {e}")
            return None

# Debug the table structure first
extractor = FireProtectionReportExtractor('/Users/austinrakowski/dev/random/firstresponse')
extractor.test_single_file('EXT Report Template (Filled).docx', 'extract_entinguishers')